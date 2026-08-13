#!/usr/bin/env python3
from datetime import date
from docx import Document
from docx.shared import Pt

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")

doc.add_heading("Atomik Acoustic Simulation Engine", level=0)
sub = doc.add_paragraph()
r = sub.add_run("UI Redesign Change Log - Graph colors.pdf to beta v1.2.3")
r.bold = True
r.font.size = Pt(14)

meta = doc.add_paragraph()
run = meta.add_run("Document date: " + date.today().isoformat() + "\n")
run.italic = True
run = meta.add_run("Scope: Brand theme, light/dark chrome, sidebar density, heatmap colours, speaker markers, plot chrome, packaging.\n")
run.italic = True
run = meta.add_run("Source reference: Graph colors.pdf + light/dark mockup screenshots provided by product owner.")
run.italic = True

para("This document summarises every major UI change implemented after Graph colors.pdf was supplied as the brand and visual source of truth. Changes apply to the native JUCE application (TwoSpeakerExplorer / Atomik Acoustic Simulation Engine). Acoustic engine behaviour is unchanged unless noted; this work is chrome, theming, and panel layout.")

doc.add_heading("1. Brand colours from Graph colors.pdf", level=1)
para("Brand tokens were centralised in Source/BrandTheme.h so light mode, dark mode, and fixed accents all read from one place.")

doc.add_heading("1.1 Fixed brand tokens", level=2)
bullet("Charcoal #231F20 - deep ink / plot floor")
bullet("White #FFFFFF - contrast / dark-mode control surfaces")
bullet("Ash #A7A9AC - secondary brand grey (reference)")
bullet("Signal Red #ED2227 - primary accent (replaced former cyan accents)")
bullet("Heatmap cool stops: #3281B9, #0A4D74, #9FB4D1")
bullet("Heatmap hot stops: #530000, #B21619, #ED2227")

doc.add_heading("1.2 Accent system", level=2)
bullet("Interactive accent (sliders, checked toggles, active pills) switched from cyan to Signal Red #ED2227.")
bullet("Checked checkboxes: red fill with white tick.")
bullet("Slider thumbs: solid red circles; filled track portion red; empty track dark grey.")
bullet("Light-mode plot titles use signal red for brand emphasis.")

doc.add_heading("2. Dual theme (Light / Dark)", level=1)
para("Theme preference is persisted in AppSettings and flips chrome only. The Rel. SPL heatmap colour map stays identical in both themes.")

doc.add_heading("2.1 Dark mode", level=2)
bullet("Near-black chrome: panels / base about #1C1C1C")
bullet("Control surfaces (buttons, combos, value boxes): white")
bullet("Ink on controls: charcoal / black so text remains readable on white pills")
bullet("Header Stats / help icons restyled for dark mockup (white Stats pill, etc.)")

doc.add_heading("2.2 Light mode", level=2)
bullet("Warm panel / app background: #EFE8E8")
bullet("Control fill inside boxes / buttons / dropdowns: #DFDEDE (user-specified)")
bullet("Solid near-black ink for labels and headers (100% opacity - no washed grey text)")
bullet("Borders: dark hairlines (about #6E6A6A) for definition on light grey boxes")
bullet("Plot canvas: light warm off-white; grid using charcoal with transparency")

doc.add_heading("3. Typography", level=1)
bullet("Primary UI font: Montserrat (Regular / Medium / SemiBold / Bold bundled under Assets)")
bullet("Numeric / digital readouts: Space Mono where appropriate (slider value boxes)")
bullet('Section headers: ALL-CAPS SemiBold (e.g. "1. FREQUENCY (Hz)", "3. SELECTED XN18", "4. SIMULATION")')
bullet("Field labels: Montserrat SemiBold, solid black, Title Case where specified (Invert Polarity, Grid Resolution, db Floor)")
bullet("Central size control: Source/UiTextConfig.h -> UiConfig::FontSize")

doc.add_heading("3.1 Sidebar size targets (late iteration)", level=2)
bullet("Section headers: 14.5 pt")
bullet("Field / checkbox labels: 14.5 pt SemiBold")
bullet("Inside-box text (combo values, buttons, numeric boxes): about 13.5-14.5 pt")
bullet("Quick Layout helper and related labels raised to match Selected XN18 crop references")
bullet("Panel remains scrollable so taller type still fits on all displays")

doc.add_heading("3.2 Where to edit fonts", level=2)
bullet("Source/UiTextConfig.h - sectionHeader, fieldLabel, fieldValue, button")
bullet("Source/ControlPanel.cpp - layout helper font in updateScaledChrome()")
bullet("Source/BrandTheme.h - LookAndFeel getComboBoxFont / getTextButtonFont / drawToggleButton")

doc.add_heading("4. SPL Heatmap colour map (Rel. SPL)", level=1)
para('ColourMaps::sevenColor implements the Graph colors.pdf ladder. Legend ticks remain 0 db to -36 db with caption "Rel. SPL".')
bullet("Top (0 db): dark maroon #530000")
bullet("Then #B21619 -> #ED2227")
bullet("Magenta bridge into cool band")
bullet("Then #3281B9 -> #0A4D74 -> charcoal #231F20 at the bottom (-36)")
bullet("Theme changes do not recolour the field map stops")

doc.add_heading("5. Speaker markers on the heatmap", level=1)
bullet("White rounded tiles with black speaker glyph (cone + wave arcs)")
bullet("Waves oriented to the right; drawing clipped/padded so waves fit")
bullet("Labels above cabinets: XN18_n")
bullet("Selected vs unselected id weight distinction preserved")

doc.add_heading("6. Left control sidebar structure and density", level=1)
para("Layout and labels were iterated against provided light-panel mockups until density, alignment, and copy matched.")

doc.add_heading("6.1 Section order", level=2)
bullet("1. FREQUENCY (Hz) - frequency dropdown")
bullet("2. XN18 Units - unit combo, + Add, Delete; Quick Layout (same plane); 1 / 2 / 3 Device buttons")
bullet("3. SELECTED XN18 - X/Y Position, Gain, Delay sliders + Invert Polarity / Reverse Orientation / Enabled")
bullet("4. SIMULATION - Grid Resolution, db Floor, contour/smoothing/measured-directivity toggles, measurement set")
bullet("5. WORKSPACE / 6. ARRAY PRESETS - collapsible (can remain collapsed)")
bullet("Reset link at foot of panel")

doc.add_heading("6.2 Control chrome", level=2)
bullet("Row layout: label left | thin slider | numeric box right (value boxes aligned)")
bullet("Tight vertical gaps (about 33 px row pitch in compact mockup measurements)")
bullet("Slider track thin; thumb diameter about 10 px red")
bullet("Unchecked checkboxes filled with #DFDEDE; checked = signal red + white tick")
bullet("Combo / button / value-box fill: #DFDEDE in light mode")
bullet("Slightly larger in-box type and value box size (52x24) for legibility")

doc.add_heading("6.3 Label copy alignment with mockups", level=2)
bullet("X Position (m), Y Position (m), Gain (db), Delay (ms)")
bullet("Invert Polarity, Reverse Orientation, Enabled")
bullet("Grid Resolution, db Floor")
bullet("Quick Layout (same plane)")

doc.add_heading("7. Collapsible sidebar", level=1)
bullet('User can hide/show the left controls with a bottom-of-sidebar button ("Hide panel" / "Show").')
bullet("When collapsed, a thin left rail remains with Show at the bottom; plot and bottom toolbar reclaim width.")
bullet("State is persisted in AppSettings (sidebarCollapsed) and restored on next launch.")
bullet("Widths scale with the responsive UI scale so behaviour stays correct on different display sizes.")

doc.add_heading("8. Header, plot chrome, and bottom bar", level=1)
bullet('App title: "Atomik Acoustic Simulation Engine"; version label (e.g. beta_v1.2).')
bullet("Header actions: Stats dropdown, Help, Preferences (gear), More overflow.")
bullet("Plot header toolbar icons restyled for theme (e.g. black icons on white buttons in dark theme).")
bullet("Bottom SAVE/EXPORT and VIEW MODE pill layouts kept compact to match v1.1+ mockups.")
bullet("Status strip shows last run / ready state.")

doc.add_heading("9. Responsive UI scale", level=1)
para("UiConfig::Scale uses a baseline window of 1340x820. Fonts and layout pixels scale uniformly within min/max factors so the app remains usable when maximised or on larger monitors without separately hand-tuning every control.")

doc.add_heading("10. Local run / packaging notes", level=1)
bullet("Windows Smart App Control could block unsigned Debug CRT builds. Debug uses release CRT + PDB; post-build Tools/SignLocal.ps1 signs with local cert CN=Atomik Local Dev.")
bullet("Trusted Publisher install of that cert was required so F5 launches succeed under Application Control.")
bullet("Release artifact delivered as Builds/Release/betav1.2.3.exe (signed copy of the Release build).")

doc.add_heading("11. Key source files touched", level=1)
bullet("Source/BrandTheme.h - palettes, accent, LookAndFeel (buttons, combos, sliders, checkboxes)")
bullet("Source/UiTextConfig.h - font sizes and layout metrics (single dial for designers)")
bullet("Source/ControlPanel.h / .cpp - sidebar sections, copy, resized density")
bullet("Source/UiChrome.h - section header chevrons / titles")
bullet("Source/ColourMaps.h - sevenColor Rel. SPL ladder")
bullet("Source/RadiationPatternComponent.cpp - speaker glyph tiles + labels")
bullet("Source/MainComponent.cpp - sidebar collapse, overall shell layout, theme sync")
bullet("Source/AppSettings.h - theme, grid, sidebarCollapsed persistence")
bullet("Tools/SignLocal.ps1 - local code signing for Application Control")

doc.add_heading("12. Before to After (summary)", level=1)
bullet("Accent: cyan-style interactive -> Signal Red #ED2227 from Graph colors.pdf")
bullet("Light controls: generic grey -> #DFDEDE boxes on #EFE8E8 panels")
bullet("Text: mixed weights / soft greys -> Montserrat, solid opaque near-black labels")
bullet("Sidebar: loose spacing / uneven hierarchy -> numbered ALL-CAPS sections, compact rows, mockup copy")
bullet("Heatmap: prior map -> Graph colors.pdf seven-stop Rel. SPL ladder")
bullet("Speakers: simple markers -> white tile + black glyph + XN18_n labels")
bullet("Sidebar: always visible -> optional collapse with bottom toggle, persisted")
bullet("Deliverable: TwoSpeakerExplorer.exe -> also packaged as betav1.2.3.exe")

doc.add_heading("13. Out of scope / unchanged", level=1)
bullet("Core acoustic computation paths remain SI; unit system preference is display-only.")
bullet("Measured polar datasets and project file format were not redesign goals of this UI pass.")
bullet("Further pixel-perfect tweaks can continue via UiTextConfig.h without structural rewrites.")

end = doc.add_paragraph()
end.add_run("End of change log").italic = True

out = r"d:\shayam gui\Atomik_UI_ChangeLog_GraphColors_to_betav1.2.3.docx"
doc.save(out)
print("Wrote", out)
